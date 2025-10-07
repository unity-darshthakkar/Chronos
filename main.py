import os
import re
import json
import time
import wave
import queue
import sys
import threading
from typing import Union
from urllib.parse import quote
from zoneinfo import ZoneInfo
import subprocess
import webbrowser
import string
import pathlib
from typing import Optional, List, Tuple, Dict, Set
import requests
import sounddevice as sd
import pyperclip
from datetime import datetime
import shutil  # for PATH checking

# NEW: for writing the Word doc
try:
    from docx import Document
except Exception:
    Document = None  # we'll warn at runtime if missing

try:
    import win32com.client as win32  # Office automation
except Exception:
    win32 = None

try:
    from send2trash import send2trash
except Exception:
    send2trash = None

try:
    import keyboard  # media keys
except Exception:
    keyboard = None

# Optional UI automation/hotkeys for Save As fallback and Alt+F4
try:
    import pyautogui
except Exception:
    pyautogui = None

from faster_whisper import WhisperModel

# ---------- Global geocoding/timezone helpers (worldwide cities, no hardcoding) ----------
from functools import lru_cache
from geopy.geocoders import Nominatim
from geopy.extra.rate_limiter import RateLimiter
from timezonefinder import TimezoneFinder

_geolocator = Nominatim(user_agent="friday-assistant/1.0", timeout=5)
_geocode = RateLimiter(_geolocator.geocode, min_delay_seconds=1)
_tzf = TimezoneFinder()

# Bubble/tray mute state
_muted = False
_last_opened_doc_path: Optional[str] = ""

@lru_cache(maxsize=1024)
def _geocode_city(city_query: str) -> Optional[Tuple[str, float, float]]:

    if not city_query:
        return None
    try:
        loc = _geocode(city_query)
        if not loc:
            return None
        name = loc.address
        return (name, float(loc.latitude), float(loc.longitude))
    except Exception:
        return None

@lru_cache(maxsize=1024)
def _geocode_city_details(city_query: str) -> Optional[Dict[str, Union[str, float]]]:
    """
    Returns dict with display_name, lat, lon, country_code (ISO-2, upper).
    """
    try:
        loc = _geocode(city_query)
        if not loc:
            return None
        addr = getattr(loc, "raw", {}).get("address", {}) or {}
        cc = (addr.get("country_code") or "").upper()
        return {
            "display_name": loc.address,
            "lat": float(loc.latitude),
            "lon": float(loc.longitude),
            "country_code": cc
        }
    except Exception:
        return None

@lru_cache(maxsize=2048)
def _latlon_to_tz(lat: float, lon: float) -> Optional[str]:
    try:
        tz = _tzf.timezone_at(lat=lat, lng=lon)
        return tz
    except Exception:
        return None


def _resolve_city_to_tz(city: str) -> Tuple[str, Optional[str]]:
    if not city:
        return ("", None)
    res = _geocode_city(city.strip())
    if not res:
        return (city.strip().title(), None)
    display_name, lat, lon = res
    tz = _latlon_to_tz(lat, lon)
    return (display_name, tz)


# ============================ Config ============================
SAMPLE_RATE = 16000
CHANNELS = 1
CHUNK_SEC = float(os.getenv("FRIDAY_CHUNK_SEC", 10.0))
MODEL_SIZE = os.getenv("WHISPER_MODEL", "base.en")         # tiny/base/small/medium/large-v3
DEVICE = os.getenv("WHISPER_DEVICE", "cpu")              # "cpu" or "cuda"
COMPUTE_TYPE = os.getenv("WHISPER_COMPUTE", "int8")      # int8/int8_float16/float16/float32

# News API keys (optional)
NEWSAPI_KEY = os.getenv("NEWSAPI_KEY", "").strip()
GNEWS_API_KEY = os.getenv("GNEWS_API_KEY", "").strip()

# Wake phrases
WAKE_PHRASES = ("hey friday", "hello friday")

# Follow-up dialog state: waiting for app name after “Which app?”
_awaiting_app_name = False

# Track opened app PIDs (best effort) and last opened
_OPENED_PIDS: Dict[str, Set[int]] = {}
_last_opened_app_key: Optional[str] = None

def _record_opened_pid(app_key: str, pid: Optional[int]):
    global _OPENED_PIDS, _last_opened_app_key
    _last_opened_app_key = app_key
    if pid is None:
        return
    _OPENED_PIDS.setdefault(app_key, set()).add(pid)

# App aliases → launched via shell PATH or explicit paths
APP_ALIASES = {
    # Office
    "word":       ["winword", "Microsoft Word"],
    "excel":      ["excel", "Microsoft Excel"],
    "powerpoint": ["powerpnt", "Microsoft PowerPoint"],
    "access":     ["msaccess", "Microsoft Access"],

    # Browsers
    "chrome":   ["chrome", "Google Chrome"],
    "edge":     ["msedge", "Microsoft Edge"],

    # Editors
    "vscode":   ["code", "Visual Studio Code"],
    "notepad":  ["notepad"],

    # Media
    "spotify":  ["spotify"],

    # Terminals / shells
    "terminal": ["wt", "Windows Terminal"],
    "cmd":      ["cmd"],
    "powershell": ["powershell"],
    "wsl":      ["wsl"],

    # System apps
    "calculator": ["calc", "calculator"],
    "explorer":   ["explorer", "file explorer"],
    "calendar":   ["calendar"],
    "browser":    ["browser"],
}


PROC_NAME_MAP = {
    "word": ["WINWORD", "WINWORD.EXE"],
    "excel": ["EXCEL", "EXCEL.EXE"],
    "powerpoint": ["POWERPNT", "POWERPNT.EXE"],
    "access": ["MSACCESS", "MSACCESS.EXE"],
    "vscode": ["Code", "Code.exe"],
    "notepad": ["notepad", "notepad.exe"],
    "chrome": ["chrome", "chrome.exe"],
    "edge": ["msedge", "msedge.exe"],
    "spotify": ["Spotify", "Spotify.exe"],
    "terminal": ["WindowsTerminal", "WindowsTerminal.exe", "wt", "OpenConsole"],
    "cmd": ["cmd", "cmd.exe"],
    "powershell": ["powershell", "powershell.exe", "pwsh", "pwsh.exe"],
    "wsl": ["wsl", "wsl.exe"],
    "explorer": ["explorer", "explorer.exe"],
    "calculator": ["CalculatorApp", "CalculatorApp.exe", "Calculator", "Calculator.exe"],
    "browser": ["chrome", "chrome.exe", "msedge", "msedge.exe"],
}

APP_PATHS = {
    # Common Office installs
    "word_paths": [
        r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\WINWORD.EXE",
        r"C:\Program Files\Microsoft Office\root\Office15\WINWORD.EXE",
        r"C:\Program Files\Microsoft Office\Office16\WINWORD.EXE",
        r"C:\Program Files\Microsoft Office\Office15\WINWORD.EXE",
    ],
    "excel_paths": [
        r"C:\Program Files\Microsoft Office\root\Office16\EXCEL.EXE",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\EXCEL.EXE",
        r"C:\Program Files\Microsoft Office\root\Office15\EXCEL.EXE",
        r"C:\Program Files\Microsoft Office\Office16\EXCEL.EXE",
        r"C:\Program Files\Microsoft Office\Office15\EXCEL.EXE",
    ],
    "powerpoint_paths": [
        r"C:\Program Files\Microsoft Office\root\Office16\POWERPNT.EXE",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\POWERPNT.EXE",
        r"C:\Program Files\Microsoft Office\root\Office15\POWERPNT.EXE",
        r"C:\Program Files\Microsoft Office\Office16\POWERPNT.EXE",
        r"C:\Program Files\Microsoft Office\Office15\POWERPNT.EXE",
    ],
    "access_paths": [
        r"C:\Program Files\Microsoft Office\root\Office16\MSACCESS.EXE",
        r"C:\Program Files (x86)\Microsoft Office\root\Office16\MSACCESS.EXE",
        r"C:\Program Files\Microsoft Office\root\Office15\MSACCESS.EXE",
        r"C:\Program Files\Microsoft Office\Office16\MSACCESS.EXE",
        r"C:\Program Files\Microsoft Office\Office15\MSACCESS.EXE",
    ],

    # Others
    "chrome":   r"C:\Program Files\Google\Chrome\Application\chrome.exe",
    "edge":     r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    "vscode":   r"%LOCALAPPDATA%\Programs\Microsoft VS Code\Code.exe",
    "notepad":  r"C:\Windows\system32\notepad.exe",
    "spotify":  r"%APPDATA%\Spotify\Spotify.exe",
    "terminal": r"%LOCALAPPDATA%\Microsoft\WindowsApps\wt.exe",
}

# LLM endpoint (OpenAI-compatible)
OPENAI_BASE_URL = os.getenv("OPENAI_BASE_URL", "http://localhost:8080")
OPENAI_API_KEY  = os.getenv("OPENAI_API_KEY", "sk-no-key")
OPENAI_MODEL    = os.getenv("OPENAI_MODEL", "mistral-7b-instruct-v0.2-q4_K_M")


# ============================ Audio I/O ============================
def record_wav_tmp(path: str, seconds: float = CHUNK_SEC, samplerate: int = SAMPLE_RATE):
    audio = sd.rec(int(seconds * samplerate), samplerate=samplerate, channels=CHANNELS, dtype='int16')
    sd.wait()
    with wave.open(path, 'wb') as wf:
        wf.setnchannels(CHANNELS)
        wf.setsampwidth(2)
        wf.setframerate(samplerate)
        wf.writeframes(audio.tobytes())
    return path


print(f"[🌀] Loading faster-whisper '{MODEL_SIZE}' on {DEVICE} ({COMPUTE_TYPE}) …")
whisper = WhisperModel(MODEL_SIZE, device=DEVICE, compute_type=COMPUTE_TYPE)

def transcribe(path: str) -> str:
    segments, _info = whisper.transcribe(path, vad_filter=True)
    text = ''.join(seg.text for seg in segments).strip()
    if text:
        print(f"[📝] You said: {text}")
    return text


# ============================ NEW: Meeting Recorder ============================
def _next_meeting_basename(dir_path: str) -> str:
    patt = re.compile(r"^meeting recording \[(\d+)\]\.wav$", re.IGNORECASE)
    max_n = 0
    for name in os.listdir(dir_path):
        m = patt.match(name)
        if m:
            try:
                n = int(m.group(1))
                max_n = max(max_n, n)
            except:
                pass
    return f"meeting recording [{max_n + 1}]"

class MeetingRecorder:
    """
    Background mic capture to WAV until stopped.
    On stop, caller can transcribe and write a .docx with the same basename.
    """
    def __init__(self, samplerate=SAMPLE_RATE, channels=CHANNELS, blocksize=2048):
        self.samplerate = samplerate
        self.channels = channels
        self.blocksize = blocksize

        self._q: "queue.Queue[bytes]" = queue.Queue()
        self._writer = None
        self._stream: Optional[sd.InputStream] = None
        self._thread: Optional[threading.Thread] = None
        self._running = False

        self.dir_path = os.getcwd()
        self.basename: Optional[str] = None
        self.audio_path: Optional[str] = None
        self.docx_path: Optional[str] = None

    @property
    def running(self) -> bool:
        return self._running

    def _callback(self, indata, frames, time_info, status):
        if status:
            pass
        self._q.put(bytes(indata))

    def start(self):
        if self._running:
            return False
        self.basename = _next_meeting_basename(self.dir_path)
        self.audio_path = os.path.join(self.dir_path, f"{self.basename}.wav")
        self.docx_path = os.path.join(self.dir_path, f"{self.basename}.docx")

        try:
            import soundfile as sf
        except Exception:
            speak("Recording needs the 'soundfile' package. Run: pip install soundfile")
            return False
        
        self._writer = sf.SoundFile(self.audio_path, mode='w',
                                    samplerate=self.samplerate,
                                    channels=self.channels, subtype='PCM_16')
        self._stream = sd.InputStream(samplerate=self.samplerate,
                                      channels=self.channels,
                                      blocksize=self.blocksize,
                                      dtype='int16',
                                      callback=self._callback)
        self._running = True
        self._thread = threading.Thread(target=self._drain_loop, daemon=True)
        self._thread.start()
        self._stream.start()
        return True

    def _drain_loop(self):
        while self._running:
            try:
                chunk = self._q.get(timeout=0.5)
                if self._writer:
                    self._writer.buffer_write(chunk, dtype='int16')
            except queue.Empty:
                continue
            except Exception:
                break

    def stop(self) -> Optional[str]:
        if not self._running:
            return None
        self._running = False
        try:
            if self._stream:
                self._stream.stop()
                self._stream.close()
        finally:
            self._stream = None

        # flush remaining
        while not self._q.empty():
            try:
                chunk = self._q.get_nowait()
                if self._writer:
                    self._writer.buffer_write(chunk, dtype='int16')
            except Exception:
                break

        try:
            if self._writer:
                self._writer.close()
        finally:
            self._writer = None

        if self._thread and self._thread.is_alive():
            self._thread.join(timeout=2.0)
        self._thread = None

        return self.audio_path

    def write_docx(self, text: str):
        if Document is None:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        doc = Document()
        doc.add_heading(self.basename or "Meeting Recording", level=1)
        meta = doc.add_paragraph()
        meta.add_run("Created: ").bold = True
        meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        doc.add_paragraph("")
        for para in (text or "").split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(self.docx_path)

meeting_rec = MeetingRecorder()


# ============================ TTS (robust, queued) ============================
import threading, queue, sys

class _TTSThread(threading.Thread):
    def __init__(self):
        super().__init__(daemon=True)
        self.q: "queue.Queue[tuple[str, threading.Event|None]]" = queue.Queue()
        self.ready = threading.Event()
        self._stop = threading.Event()

    def run(self):
        # Initialize pyttsx3 INSIDE the thread for stability on Windows/SAPI
        import pyttsx3
        try:
            engine = pyttsx3.init()
            engine.setProperty("volume", 1.0)
        except Exception as e:
            print(f"[TTS] init failed: {e}", file=sys.stderr)
            return
        self.ready.set()

        while not self._stop.is_set():
            try:
                text, done = self.q.get(timeout=0.2)
            except queue.Empty:
                continue

            if text is None:  # shutdown signal
                break

            # Speak
            try:
                print(f"[🗣️ ] {text}")
                engine.setProperty("volume", 1.0)
                engine.say(text)
                engine.runAndWait()
            except Exception as e:
                # Try one re-init if the engine dies mid-run
                try:
                    engine = pyttsx3.init()
                    engine.setProperty("volume", 1.0)
                    engine.say(text)
                    engine.runAndWait()
                except Exception as e2:
                    print(f"[TTS] Error speaking: {e2}", file=sys.stderr)

            if done is not None:
                done.set()

        # best-effort cleanup
        try:
            engine.stop()
        except Exception:
            pass

# global singleton
_tts_thread = _TTSThread()
_tts_thread.start()
_tts_thread.ready.wait(timeout=5.0)

def speak(text: str, *, block: bool = False):
    """
    Queue text to speak. By default it's non-blocking.
    Use speak('...', block=True) if you need to wait until it finishes.
    """
    if not text:
        return
    done = threading.Event() if block else None
    try:
        _tts_thread.q.put((str(text), done), timeout=0.5)
        if done is not None:
            done.wait()
    except Exception as e:
        # Fallback: last resort print if queue/thread is unavailable
        print(f"[🗣️ ] {text} (no TTS: {e})")


def speak_flush():
    """Block until everything queued so far has finished speaking."""
    done = threading.Event()
    try:
        _tts_thread.q.put(("", done), timeout=0.5)  # empty utterance as a barrier
        done.wait()
    except Exception:
        pass

def shutdown_tts():
    """Call once on program exit if you want a clean shutdown."""
    try:
        _tts_thread.q.put((None, None))  # signal thread to stop
    except Exception:
        pass

# ============================ LLM helpers ============================
INTENT_SYSTEM = (
    "You are an intent classifier for a voice assistant. The user's input starts AFTER the wake phrase.\n"
    "Choose ONE intent and return JSON ONLY (no prose, no markdown).\n"
    "Intents:\n"
    "- open_app:            {\\\"app\\\": string}\n"
    "- close_app:           {\\\"app\\\": string?}\n"
    "- open_terminal:       {}\n"
    "- open_powershell:     {}\n"
    "- open_cmd:            {}\n"
    "- open_wsl:            {}\n"
    "- open_browser:        {\\\"url\\\": string?}\n"
    "- browser_search:      {\\\"query\\\": string, \\\"new_tab\\\": boolean?}\n"
    "- web_search:          {\\\"query\\\": string}\n"
    "- get_time:            {\\\"city\\\": string}\n"
    "- get_weather:         {\\\"city\\\": string}\n"
    "- open_document:       {\"path\": string}\n"
    "- edit_document:       {\"instruction\": string}\n"
    "- get_news:            {\\\"city\\\": string}\n"
    "- get_local_time:      {}\n"
    "- clear_folder:        {\\\"path\\\": string, \\\"permanent\\\": boolean?}\n"
    "- delete_path:         {\\\"path\\\": string, \\\"permanent\\\": boolean?}\n"
    "- restore_recycle_bin: {\\\"name_contains\\\": string?}\n"
    "- write_essay:         {\\\"topic\\\": string, \\\"length_words\\\": integer?}\n"
    "- generate_code:       {\\\"language\\\": string?, \\\"prompt\\\": string, \\\"open_in\\\": string?}\n"
    "- spotify_control:     {\\\"action\\\": string, \\\"playlist_url\\\": string?}\n"
    "- ask:                 {\\\"question\\\": string}\n"
    "- multi_action:        {\\\"steps\\\": [{\\\"intent\\\": string, \\\"args\\\": object}]}\n"
    "- save_document:      {\\\"app\\\": string, \\\"path\\\": string}\n"
    "- meeting_record_start: {}\n"
    "- meeting_record_stop:  {}\n"
    "Return strictly: {\\\"intent\\\":..., \\\"args\\\":{...}}\n"
)

def _openai_chat_url():
    base = OPENAI_BASE_URL.rstrip('/')
    # Always return the OpenAI-style endpoint
    if base.endswith("/v1"):
        return f"{base}/chat/completions"
    return f"{base}/v1/chat/completions"

def call_llm(messages, *, system=None, temperature=0.7, max_tokens=4096) -> str:
    import requests

    # ---- Attempt A: OpenAI-compatible
    try:
        url = _openai_chat_url()
        headers = {
            "Authorization": f"Bearer {OPENAI_API_KEY}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": OPENAI_MODEL,
            "messages": ([{"role": "system", "content": system}] if system else []) + messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": False
        }
        r = requests.post(url, headers=headers, json=payload, timeout=60)
        if r.status_code // 100 == 2:
            data = r.json()
            return data["choices"][0]["message"]["content"]
        # If it's a 404, fall through to Attempt B without raising
        if r.status_code != 404:
            raise RuntimeError(f"LLM HTTP {r.status_code} at {url}")
    except requests.RequestException:
        pass
    except Exception:
        pass

    # ---- Attempt B: text-generation-webui style
    # IMPORTANT: do NOT carry '/v1' into this base
    api_base = OPENAI_BASE_URL.rstrip('/')
    if api_base.endswith("/v1"):
        api_base = api_base[:-3]  # strip the trailing '/v1'

    try:
        url2 = f"{api_base}/api/v1/chat"
        full_prompt = ""
        if system:
            full_prompt += f"[System]\n{system}\n\n"
        for m in messages:
            role = m.get("role", "user").capitalize()
            full_prompt += f"[{role}]\n{m.get('content','')}\n\n"
        payload2 = {
            "user_input": full_prompt.strip(),
            "mode": "chat",
            "max_new_tokens": max_tokens,
            "temperature": temperature,
            "history": {"internal": [], "visible": []}
        }
        r2 = requests.post(url2, json=payload2, timeout=60)
        r2.raise_for_status()
        data2 = r2.json()
        if "results" in data2 and data2["results"]:
            txt = data2["results"][0].get("text", "")
            if txt:
                return txt
        raise RuntimeError(f"Unexpected response from {url2}")
    except requests.RequestException as e:
        raise RuntimeError(f"LLM connection failed: {e}")


# ============================ Save path helpers & plan augmenter ============================
_KNOWN_DIRS = {
    "desktop": pathlib.Path.home() / "Desktop",
    "documents": pathlib.Path.home() / "Documents",
    "downloads": pathlib.Path.home() / "Downloads",
}

def _clean_filename(name: str) -> str:
    # basic filename sanitizer for voice-captured names
    name = name.strip().strip('."\': ')
    name = re.sub(r"[\\/:*?\"<>|]", "_", name)
    return name or "Document"

def _resolve_spoken_location(loc: str) -> Optional[pathlib.Path]:
    if not loc:
        return None
    loc_low = loc.lower().strip()
    # try exact known keys
    if loc_low in _KNOWN_DIRS:
        return _KNOWN_DIRS[loc_low]
    # allow "my desktop", "on desktop", "the downloads folder", etc.
    for key, p in _KNOWN_DIRS.items():
        if key in loc_low:
            return p
    # absolute Windows path spoken like "C slash Users slash ...": not great for voice; skip here
    return None

def _extract_save_request(utterance: str) -> Optional[dict]:
    """
    Parse phrases like:
      - save it to my desktop as test
      - save to documents named meeting notes
      - save in downloads as report v2
      - save as C:\\Users\\Me\\Work\\demo (if the user actually says a full path)
    Returns: {app: 'word', path: 'C:\\...\\name.docx'}  (only for Word for now)
    """
    u = utterance.strip()
    low = u.lower()

    # 1) Absolute path explicitly mentioned (rare by voice, useful in text)
    m_abs = re.search(r"\bsave\b.*?\b(?:as|to)\b\s*([A-Za-z]:\\[^\s]+)", u)  # keep original case for path
    if m_abs:
        raw = m_abs.group(1).strip().strip('"')
        base = pathlib.Path(raw)
        if base.suffix.lower() not in (".docx", ".doc"):
            base = base.with_suffix(".docx")
        try:
            base.parent.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
        return {"app": "word", "path": str(base)}

    # 2) “save ... (to|in|on) <folder> (as|named) <name>”
    m = re.search(r"\bsave\b.*?\b(?:to|in|on)\b\s+(.*?)\s+\b(?:as|named)\b\s+([^,\.]+)", low)
    folder_hint = None
    name = None
    if m:
        folder_hint = m.group(1).strip()
        name = _clean_filename(m.group(2))
    else:
        # 3) “save ... (as|named) <name> (to|in|on) <folder>”
        m2 = re.search(r"\bsave\b.*?\b(?:as|named)\b\s+([^,\.]+?)\s+\b(?:to|in|on)\b\s+(.*?)(?:[\,\.]|$)", low)
        if m2:
            name = _clean_filename(m2.group(1))
            folder_hint = m2.group(2).strip()

    if name:
        folder = _resolve_spoken_location(folder_hint or "")
        if folder:
            full = (folder / name).with_suffix(".docx")
            try:
                full.parent.mkdir(parents=True, exist_ok=True)
            except Exception:
                pass
            return {"app": "word", "path": str(full)}

    # 4) “save ... as <name>” with no folder -> default to Desktop
    m3 = re.search(r"\bsave\b.*?\b(?:as|named)\b\s+([^,\.]+)", low)
    if m3:
        name = _clean_filename(m3.group(1))
        full = (_KNOWN_DIRS["desktop"] / name).with_suffix(".docx")
        try:
            full.parent.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
        return {"app": "word", "path": str(full)}

    return None

def _augment_plan_with_save(utterance: str, steps: List[dict]) -> List[dict]:
    """
    If the user asked to save, inject a save_document step BEFORE closing the app.
    Only targets Word for now (easy to extend later).
    """
    save_req = _extract_save_request(utterance)
    if not save_req:
        return steps

    out = []
    inserted = False
    for i, s in enumerate(steps):
        # insert save right before the first close of 'word'
        if (not inserted and s.get("intent") == "close_app" and (s.get("args", {}).get("app", "") in ("word", ""))):
            out.append({"intent": "save_document", "args": save_req})
            inserted = True
        out.append(s)

    # If user never said "close word", still perform save after writing
    if not inserted:
        # place it after write_essay if present, else append
        placed = False
        tmp = []
        for s in out:
            tmp.append(s)
            if not placed and s.get("intent") == "write_essay":
                tmp.append({"intent": "save_document", "args": save_req})
                placed = True
        out = tmp if placed else (out + [{"intent": "save_document", "args": save_req}])

    return out


# ============================ Multi-action planner & executor ============================
_MULTI_SYSTEM = (
    "You are an intent planner for a voice assistant. Given a single user command that may contain multiple actions, "
    "produce a JSON array of steps where each step is one of these intents with args (same schema as INTENT_SYSTEM):\n"
    "- open_app:            {\"app\": string}\n"
    "- close_app:           {\"app\": string?}\n"
    "- open_terminal:       {}\n"
    "- open_powershell:     {}\n"
    "- open_cmd:            {}\n"
    "- open_wsl:            {}\n"
    "- open_browser:        {\"url\": string?}\n"
    "- open_document:       {\"path\": string}\n"
    "- edit_document:       {\"instruction\": string}\n"
    "- browser_search:      {\"query\": string, \"new_tab\": boolean?}\n"
    "- get_time:            {\"city\": string}\n"
    "- get_weather:         {\"city\": string}\n"
    "- save_document:       {\"app\": string, \"path\": string}\n"
    "- get_news:            {\"city\": string}\n"
    "- get_local_time:      {}\n"
    "- write_essay:         {\"topic\": string, \"length_words\": integer?}\n"
    "- spotify_control:     {\"action\": string, \"playlist_url\": string?}\n"
    "- generate_code:       {\"language\": string?, \"prompt\": string, \"open_in\": string?}\n"
    "Rules:\n"
    "1) Return JSON ONLY: {\"steps\":[...]} .\n"
    "2) Use concise args. Infer reasonable defaults.\n"
    "3) Keep step order matching the user's phrasing.\n"
)

def _expand_and_exists(p: str) -> Optional[str]:
    if not p:
        return None
    p = os.path.expandvars(p).strip().strip('"')
    if os.path.exists(p):
        return os.path.abspath(p)
    return None

def word_open_existing(path: str) -> bool:
    """Open a .docx/.doc in Word and bring it to front."""
    if win32 is None:
        return False
    fp = _expand_and_exists(path)
    if not fp:
        speak("That file path doesn't exist.")
        return False
    try:
        app = _word_app() or win32.gencache.EnsureDispatch('Word.Application')
        app.Visible = True
        app.Documents.Open(fp)
        speak(f"Opened {os.path.basename(fp)} in Word.")
        return True
    except Exception as e:
        print(f"[WORD OPEN] {e}")
        return False

def open_document(path: str) -> bool:
    """Open any document by path. Uses Word for .doc/.docx, Notepad for .txt, else default app."""
    global _last_opened_doc_path
    fp = _expand_and_exists(path)
    if not fp:
        speak("I couldn't find that file.")
        return False

    ext = os.path.splitext(fp)[1].lower()
    try:
        if ext in (".docx", ".doc"):
            ok = word_open_existing(fp)
            if ok:
                _last_opened_doc_path = fp
                _record_opened_pid("word", None)
            return ok
        elif ext in (".txt", ".md", ".csv", ".log"):
            # simple text types -> Notepad
            p = subprocess.Popen(f'notepad "{fp}"', shell=True)
            speak(f"Opening {os.path.basename(fp)} in Notepad.")
            _record_opened_pid("notepad", getattr(p, "pid", None))
            _last_opened_doc_path = fp
            return True
        else:
            # Everything else -> system default
            os.startfile(fp)  # Windows default handler
            speak(f"Opening {os.path.basename(fp)}.")
            _last_opened_doc_path = fp
            _record_opened_pid("browser", None)  # or skip if not a browser
            return True
    except Exception as e:
        speak(f"Couldn't open the file: {e}")
        return False
    

def word_get_active_text() -> Optional[str]:
    if win32 is None:
        return None
    try:
        app = _word_app()
        if not app or app.Documents.Count == 0:
            return None
        return app.ActiveDocument.Content.Text
    except Exception:
        return None

def word_set_active_text(new_text: str) -> bool:
    """Replace whole doc content (keeps file, loses rich formatting)."""
    if win32 is None:
        return False
    try:
        app = _word_app()
        if not app or app.Documents.Count == 0:
            return False
        rng = app.ActiveDocument.Content
        rng.Text = new_text
        return True
    except Exception as e:
        print(f"[WORD SET] {e}")
        return False

def word_append_text(text: str) -> bool:
    if win32 is None:
        return False
    try:
        app = _word_app()
        if not app or app.Documents.Count == 0:
            return False
        app.Selection.EndKey(Unit=6)  # wdStory
        app.Selection.TypeParagraph()
        app.Selection.TypeText(text)
        return True
    except Exception:
        return False

def word_replace_all(find_text: str, replace_with: str) -> bool:
    """Use Word's Find/Replace across entire doc."""
    if win32 is None:
        return False
    try:
        app = _word_app()
        if not app or app.Documents.Count == 0:
            return False
        rng = app.ActiveDocument.Content
        find = rng.Find
        # wd constants: Forward=1, Wrap=1 (wdFindContinue), Replace=2 (wdReplaceAll)
        return bool(find.Execute(FindText=find_text, ReplaceWith=replace_with, Replace=2, Forward=True, Wrap=1))
    except Exception:
        return False

def _classify_edit_instruction(instr: str) -> Dict[str, str]:
    """
    Returns {'mode': 'add|update|delete|replace', 'find': ..., 'repl': ...}
    """
    s = (instr or "").strip()
    low = s.lower()

    # replace "X" with "Y"
    m = re.search(r'replace\s+"([^"]+)"\s+with\s+"([^"]+)"', low, re.I)
    if m:
        return {"mode": "replace", "find": m.group(1), "repl": m.group(2)}

    # delete "X"
    m = re.search(r'delete\s+"([^"]+)"', low, re.I) or re.search(r'remove\s+"([^"]+)"', low, re.I)
    if m:
        return {"mode": "delete", "find": m.group(1)}

    # add / append / insert
    if any(k in low for k in ("add ", "append ", "insert ")):
        return {"mode": "add"}

    # update / edit / rewrite / revise / modify / replace without quotes
    if any(k in low for k in ("update", "edit", "rewrite", "revise", "modify")):
        return {"mode": "update"}

    # default: update
    return {"mode": "update"}

def edit_active_document(instruction: str) -> bool:
    """
    Applies 'instruction' to the CURRENTLY OPEN document.
    - Word: uses COM (best).
    - Notepad/.txt: if _last_opened_doc_path points to a text file, we load/modify/save.
    """
    # Prefer Word if available and a document is open
    if win32 is not None:
        text = word_get_active_text()
        if text is not None:
            cls = _classify_edit_instruction(instruction)
            mode = cls["mode"]

            if mode == "replace" and cls.get("find") is not None:
                ok = word_replace_all(cls["find"], cls.get("repl", ""))
                speak("Replaced the text.") if ok else speak("I couldn't find that text to replace.")
                return ok

            if mode == "delete":
                target = cls.get("find")
                if target:
                    ok = word_replace_all(target, "")
                    speak("Deleted the text.") if ok else speak("I couldn't find that text to delete.")
                    return ok
                # No explicit target: fall back to LLM rewrite that removes described content
                mode = "update"

            if mode == "add":
                # Ask LLM to draft the new section based on instruction and existing context
                prompt = f"You are editing a document. Based on this request:\n\n{instruction}\n\nWrite ONLY the new content to add. Do not repeat existing text."
                try:
                    addition = call_llm([{"role":"user","content":prompt}],
                                        system="You return only the text to append. No headings unless asked.",
                                        temperature=0.6, max_tokens=800).strip()
                except Exception as e:
                    speak(f"LLM error: {e}")
                    return False
                if not addition:
                    speak("I didn't generate anything to add.")
                    return False
                ok = word_append_text(addition)
                speak("Added the new content.") if ok else speak("Couldn't append to the document.")
                return ok

            if mode == "update":
                # LLM transforms whole document according to instruction
                sys_prompt = ("You edit documents. Apply the user's instruction to the document text faithfully. "
                              "Return the FULL revised document text. Keep it coherent and concise.")
                user = f"Instruction: {instruction}\n\n---\nCurrent document:\n{text}\n---\nReturn the revised document."
                try:
                    revised = call_llm([{"role":"user","content":user}],
                                       system=sys_prompt, temperature=0.5, max_tokens=4096).strip()
                except Exception as e:
                    speak(f"LLM error: {e}")
                    return False
                if not revised:
                    speak("No changes were made.")
                    return False
                ok = word_set_active_text(revised)
                speak("Updated the document.") if ok else speak("Couldn't update the document.")
                return ok

            speak("I didn't understand how to edit that.")
            return False

    # Fallback: text files via path (Notepad has no API; we edit on disk)
    fp = _last_opened_doc_path or ""
    if fp and os.path.splitext(fp)[1].lower() in (".txt", ".md", ".csv", ".log"):
        try:
            with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                cur = f.read()
        except Exception as e:
            speak(f"Couldn't read the file: {e}")
            return False

        cls = _classify_edit_instruction(instruction)
        mode = cls["mode"]
        out = cur

        if mode == "replace" and cls.get("find") is not None:
            out = cur.replace(cls["find"], cls.get("repl", ""))
        elif mode == "delete":
            target = cls.get("find")
            if target:
                out = cur.replace(target, "")
            else:
                # LLM rewrite to remove described content
                try:
                    out = call_llm([{"role":"user","content":f"Instruction: {instruction}\n\nDocument:\n{cur}\n\nReturn the revised document text."}],
                                   system="You edit plain text per instruction and return full result.",
                                   temperature=0.5, max_tokens=4096).strip()
                except Exception as e:
                    speak(f"LLM error: {e}")
                    return False
        elif mode == "add":
            try:
                addition = call_llm([{"role":"user","content":f"Based on: {instruction}\n\nWrite only the new content to append."}],
                                    system="Return only new content.", temperature=0.6, max_tokens=800).strip()
            except Exception as e:
                speak(f"LLM error: {e}")
                return False
            out = cur + ("\n\n" if cur and addition else "") + addition
        else:  # update/rewrite
            try:
                out = call_llm([{"role":"user","content":f"Instruction: {instruction}\n\nDocument:\n{cur}\n\nReturn the revised document text."}],
                               system="You edit plain text per instruction and return full result.",
                               temperature=0.5, max_tokens=4096).strip()
            except Exception as e:
                speak(f"LLM error: {e}")
                return False

        try:
            with open(fp, "w", encoding="utf-8") as f:
                f.write(out)
            speak("Updated the file.")
            return True
        except Exception as e:
            speak(f"Couldn't write changes: {e}")
            return False

    speak("Open a document first, then tell me how to edit it.")
    return False


def _planner_llm(utterance: str) -> Optional[dict]:
    try:
        out = call_llm(
            [{"role": "user", "content": utterance}],
            system=_MULTI_SYSTEM,
            temperature=0.2,
            max_tokens=512
        )
        m = re.search(r"\{[\s\S]*\}$", out.strip())
        if not m: 
            return None
        plan = json.loads(m.group(0))
        if isinstance(plan, dict) and isinstance(plan.get("steps"), list):
            return plan
    except Exception:
        pass
    return None

def _fallback_plan(utterance: str) -> Optional[dict]:
    """
    Heuristic fallback for the demo phrase:
    'open a word doc, write a 200 word essay on ai, save it my desktop using the name test and close word'
    """
    low = utterance.lower()
    steps = []

    # Open Word
    if re.search(r"\bopen\b.*\b(word|ms word|winword)\b", low):
        steps.append({"intent": "open_app", "args": {"app": "word"}})

    # Write essay (extract ~N words & topic)
    m = re.search(r"write (?:a|an)?\s*(\d{2,4})\s*word\s+essay\s+on\s+(.+?)(?:,| and | then | save | close|$)", low)
    if m:
        n = int(m.group(1))
        topic = m.group(2).strip(" .")
        steps.append({"intent": "write_essay", "args": {"topic": topic, "length_words": n}})
    else:
        m2 = re.search(r"write (?:a|an)?\s*.*?essay\s+on\s+(.+?)(?:,| and | then | save | close|$)", low)
        if m2:
            topic = m2.group(1).strip(" .")
            steps.append({"intent": "write_essay", "args": {"topic": topic, "length_words": 150}})

    # Save step (if requested)
    save_req = _extract_save_request(utterance)
    if save_req:
        steps.append({"intent": "save_document", "args": save_req})


    # Close word
    if re.search(r"\bclose\b.*\bword\b", low):
        steps.append({"intent": "close_app", "args": {"app": "word"}})

    return {"steps": steps} if steps else None

def _resolve_desktop_save_path(name: str, ext: str = ".docx") -> str:
    base = (pathlib.Path.home() / "Desktop" / name).with_suffix(ext)
    return str(base)

def _execute_step(step: dict):
    intent = (step or {}).get("intent")
    args = (step or {}).get("args", {}) or {}

    # Special placeholder: save active Word to Desktop with given name.
    if intent == "close_app" and isinstance(args.get("app"), str) and args["app"].startswith("__word_save_to_desktop__"):
        if win32 is None:
            speak("I can't save in Word without Office automation installed.")
            return
        name = args["app"].replace("__word_save_to_desktop__", "").strip() or "Document"
        full = _resolve_desktop_save_path(name, ".docx")
        if word_save_active_as(full):
            speak(f"Saved to Desktop as {pathlib.Path(full).name}.")
        else:
            speak("Couldn't save the Word document automatically. I opened Save As if possible.")
        return
    
    if intent == "save_document":
        app = (args.get("app") or "word").lower()
        path = args.get("path")
        if app == "word" and path:
            ok = word_save_active_as(path)
            if ok:
                speak(f"Saved as {pathlib.Path(path).name}.")
            else:
                speak("Couldn't save the document automatically. I opened Save As if possible.")
        else:
            speak("I can only auto-save Word documents right now.")
        return

    # Normal routing to your existing handlers:
    if intent == "open_app":           open_app(args.get("app","")); return
    if intent == "close_app":          close_app(args.get("app","")); return
    if intent == "open_terminal":      open_terminal(); return
    if intent == "open_powershell":    open_powershell(); return
    if intent == "open_cmd":           open_cmd(); return
    if intent == "open_wsl":           open_wsl(); return
    if intent == "open_browser":       open_browser(args.get("url")); return
    if intent == "browser_search":     browser_search(args.get("query",""), bool(args.get("new_tab", False))); return
    if intent == "get_time":           open_city_time(args.get("city","")); return
    if intent == "get_weather":        open_city_weather(args.get("city","")); return
    if intent == "get_news":           open_city_news(args.get("city","")); return
    if intent == "get_local_time":     speak(f"The current time is {get_local_time_str()}."); return
    if intent == "open_document":   open_document(args.get("path","")); return
    if intent == "edit_document":   edit_active_document(args.get("instruction","")); return
    if intent == "write_essay":        do_write_essay(args.get("topic",""), args.get("length_words")); return
    if intent == "spotify_control":    spotify_control(args.get("action",""), args.get("playlist_url")); return
    if intent == "generate_code":      generate_code(args.get("language"), args.get("prompt",""), args.get("open_in","vscode")); return
    # Unknown step
    speak(f"I don't know how to do '{intent}' yet.")

def execute_plan(steps: List[dict]):
    for i, step in enumerate(steps, 1):
        try:
            _execute_step(step)
            # Brief pacing to avoid UI races (Word/Notepad focus etc.)
            time.sleep(0.4)
        except SystemExit:
            raise
        except Exception as e:
            speak(f"Step {i} failed: {e}")


# ============================ Web helpers (time, browser) ============================
def web_get_time_for_city(city: str) -> Optional[Tuple[str, str]]:
    """
    Resolve any city worldwide to a timezone and get current local time via timeapi.io.
    Returns (formatted_time_12h, tzname) or None.
    """
    display_name, tzname = _resolve_city_to_tz(city)
    if not tzname:
        return None
    try:
        url = f"https://timeapi.io/api/Time/current/zone?timeZone={quote(tzname, safe='')}"
        resp = requests.get(url, timeout=8)
        if resp.status_code == 200:
            data = resp.json()
            hour = data.get("hour"); minute = data.get("minute")
            ampm = "AM" if (hour is not None and hour < 12) else "PM"
            if hour is not None and minute is not None:
                h12 = hour % 12 or 12
                tstr = f"{h12:02d}:{minute:02d} {ampm}"
                return (tstr, tzname)
    except Exception:
        pass
    return None

def get_local_time_str() -> str:
    now = datetime.now()
    ampm = "AM" if now.hour < 12 else "PM"
    h12 = now.hour % 12 or 12
    return f"{h12:02d}:{now.minute:02d} {ampm}"


# ---- simple default-browser opener ----
class BrowserController:
    def __init__(self):
        self._owned = False

    def ensure(self):
        return False

    def open_url(self, url: str, new_tab: bool = False):
        try:
            if os.name == 'nt':
                subprocess.Popen(f'start "" "{url}"', shell=True)
            else:
                webbrowser.open(url, new=new_tab)
        except Exception:
            try:
                webbrowser.open(url, new=new_tab)
            except Exception:
                pass

    def search(self, query: str, new_tab: bool = False):
        url = f"https://www.google.com/search?q={quote(query, safe='')}"
        self.open_url(url, new_tab=new_tab)

    def close(self):
        pass

_browser = BrowserController()

def open_web_shortcut(name: str):
    n = name.lower().strip()
    if n in ("gmail", "google mail", "mail gmail"):
        _browser.open_url("https://mail.google.com/"); speak("Opening Gmail."); _record_opened_pid("browser", None); return True
    if n in ("outlook", "outlook mail", "ms outlook"):
        _browser.open_url("https://outlook.live.com/mail/"); speak("Opening Outlook."); _record_opened_pid("browser", None); return True
    if n in ("google calendar", "gcal", "calendar google"):
        _browser.open_url("https://calendar.google.com/"); speak("Opening Google Calendar."); _record_opened_pid("browser", None); return True
    return False


# ============================ NEW: Weather API (Open-Meteo) ============================
_WEATHER_CODE_MAP = {
    0: "clear sky", 1: "mainly clear", 2: "partly cloudy", 3: "overcast",
    45: "fog", 48: "depositing rime fog",
    51: "light drizzle", 53: "moderate drizzle", 55: "dense drizzle",
    56: "freezing drizzle", 57: "freezing drizzle",
    61: "light rain", 63: "moderate rain", 65: "heavy rain",
    66: "freezing rain", 67: "freezing rain",
    71: "light snow", 73: "moderate snow", 75: "heavy snow",
    77: "snow grains",
    80: "light showers", 81: "moderate showers", 82: "violent showers",
    85: "snow showers", 86: "heavy snow showers",
    95: "thunderstorm", 96: "thunderstorm with hail", 99: "severe thunderstorm with hail"
}

def _use_fahrenheit(country_code: str) -> bool:
    # Imperial-leaning countries; expand if you want
    return country_code.upper() in {"US", "BS", "BZ", "KY", "PW"}

def _fetch_weather_for(lat: float, lon: float, use_f: bool) -> Optional[dict]:
    try:
        params = {
            "latitude": lat,
            "longitude": lon,
            "current": "temperature_2m,apparent_temperature,relative_humidity_2m,weather_code,wind_speed_10m,precipitation",
            "temperature_unit": "fahrenheit" if use_f else "celsius",
            "windspeed_unit": "mph" if use_f else "kmh",
            "precipitation_unit": "inch" if use_f else "mm",
            "timezone": "auto",
        }
        r = requests.get("https://api.open-meteo.com/v1/forecast", params=params, timeout=6)
        if r.status_code == 200:
            data = r.json()
            return data.get("current")
    except Exception:
        pass
    return None

def _fmt_weather_phrase(cur: dict, use_f: bool) -> str:
    if not cur:
        return ""
    t = cur.get("temperature_2m")
    at = cur.get("apparent_temperature")
    hum = cur.get("relative_humidity_2m")
    w = cur.get("wind_speed_10m")
    code = int(cur.get("weather_code") or 0)
    cond = _WEATHER_CODE_MAP.get(code, "unknown conditions")
    unit_t = "°F" if use_f else "°C"
    unit_w = "mph" if use_f else "km/h"
    parts = []
    if t is not None:
        parts.append(f"{t:.0f}{unit_t}")
    if at is not None:
        parts.append(f"feels like {at:.0f}{unit_t}")
    head = ", ".join(parts) if parts else cond
    tail = []
    if hum is not None:
        tail.append(f"humidity {hum}%")
    if w is not None:
        tail.append(f"wind {w:.0f} {unit_w}")
    tail_str = "; " + ", ".join(tail) if tail else ""
    return f"{head}, {cond}{tail_str}."

# ============================ NEW: City News (NewsAPI / GNews) ============================
def _fetch_news_newsapi(city_q: str) -> Optional[List[Dict[str, str]]]:
    if not NEWSAPI_KEY:
        return None
    try:
        params = {
            "q": city_q,
            "searchIn": "title,description",
            "language": "en",
            "sortBy": "publishedAt",
            "pageSize": 5,
            "apiKey": NEWSAPI_KEY
        }
        r = requests.get("https://newsapi.org/v2/everything", params=params, timeout=6)
        if r.status_code == 200:
            data = r.json()
            arts = data.get("articles") or []
            out = []
            for a in arts:
                title = (a.get("title") or "").strip()
                url = (a.get("url") or "").strip()
                if title:
                    out.append({"title": title, "url": url})
            return out[:5]
    except Exception:
        pass
    return None

def _fetch_news_gnews(city_q: str) -> Optional[List[Dict[str, str]]]:
    if not GNEWS_API_KEY:
        return None
    try:
        params = {
            "q": city_q,
            "lang": "en",
            "max": 5,
            "apikey": GNEWS_API_KEY
        }
        r = requests.get("https://gnews.io/api/v4/search", params=params, timeout=6)
        if r.status_code == 200:
            data = r.json()
            arts = data.get("articles") or []
            out = []
            for a in arts:
                title = (a.get("title") or "").strip()
                url = (a.get("url") or "").strip()
                if title:
                    out.append({"title": title, "url": url})
            return out[:5]
    except Exception:
        pass
    return None

def _get_city_news_list(city_display: str) -> Optional[List[Dict[str, str]]]:
    # Try NewsAPI, then GNews
    res = _fetch_news_newsapi(city_display)
    if res:
        return res
    res = _fetch_news_gnews(city_display)
    if res:
        return res
    return None


# ============================ City searches (time / weather / news) ============================
def open_city_time(city: str, open_in_browser: bool = False, new_tab: bool = False) -> bool:
    display_name, tzname = _resolve_city_to_tz(city)
    if tzname:
        info = web_get_time_for_city(city)
        if info:
            tstr, tz = info
            speak(f"The time in {display_name} is {tstr} ({tz}).")
        else:
            speak(f"Sorry, I couldn't get the time for {display_name} right now.")
    else:
        speak(f"Sorry, I couldn't find that location.")
    if open_in_browser:
        _browser.search(f"time in {display_name}", new_tab=new_tab)
        _record_opened_pid("browser", None)
    return True

def open_city_weather(city: str, open_in_browser: bool = False, new_tab: bool = False) -> bool:
    details = _geocode_city_details(city)
    if not details:
        speak("Sorry, I couldn't find that location.")
        return True
    display = details["display_name"]
    lat, lon = details["lat"], details["lon"]
    use_f = _use_fahrenheit(details["country_code"])
    cur = _fetch_weather_for(lat, lon, use_f)
    if cur:
        phrase = _fmt_weather_phrase(cur, use_f)
        speak(f"Weather in {display}: {phrase}")
    else:
        speak(f"Sorry, I couldn't fetch weather for {display} right now.")
    if open_in_browser:
        _browser.search(f"weather in {display}", new_tab=new_tab)
        _record_opened_pid("browser", None)
    return True

def open_city_news(city: str, open_in_browser: bool = False, new_tab: bool = False) -> bool:
    details = _geocode_city_details(city)
    if not details:
        speak("Sorry, I couldn't find that location.")
        return True
    display = details["display_name"]
    headlines = _get_city_news_list(display)
    if headlines:
        # Read top 3 quickly
        tops = headlines[:3]
        speak(f"Top news in {display}: " + "; ".join(h['title'] for h in tops) + ".")
    else:
        if not (NEWSAPI_KEY or GNEWS_API_KEY):
            speak(f"I couldn't access a news API. Set NEWSAPI_KEY or GNEWS_API_KEY to enable headlines. Opening browser.")
        else:
            speak(f"Sorry, I couldn't fetch headlines for {display} right now.")
    if open_in_browser or not headlines:
        _browser.search(f"{display} news", new_tab=new_tab)
        _record_opened_pid("browser", None)
    return True


# ============================ Close helpers ============================
def _ps_run(script: str) -> int:
    try:
        completed = subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            check=False
        )
        return completed.returncode
    except Exception:
        return -1

def _ps_close_main_window_by_pids(pids: List[int]) -> bool:
    if not pids:
        return False
    try:
        script = f"Get-Process -Id {','.join(str(p) for p in pids)} -ErrorAction SilentlyContinue | ForEach-Object {{ $_.CloseMainWindow() | Out-Null }}"
        rc = subprocess.call(
            ["powershell", "-NoProfile", "-Command", script],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, shell=False
        )
        return rc == 0
    except Exception:
        return False

def _ps_close_main_window_by_names(names: List[str]) -> bool:
    if not names:
        return False
    try:
        quoted = ",".join(f"'{n}'" for n in names)
        script = f"Get-Process -Name {quoted} -ErrorAction SilentlyContinue | ForEach-Object {{ $_.CloseMainWindow() | Out-Null }}"
        rc = subprocess.call(
            ["powershell", "-NoProfile", "-Command", script],
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, shell=False
        )
        return rc == 0
    except Exception:
        return False

# ---- generic taskkill (with /T: kills process tree) ----
def _taskkill_by_pids(pids: List[int], force: bool = False) -> bool:
    if not pids:
        return False
    ok = False
    for pid in pids:
        try:
            args = ["taskkill", "/PID", str(pid), "/T"]
            if force:
                args.append("/F")
            rc = subprocess.call(args, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            ok = ok or (rc == 0)
        except Exception:
            pass
    return ok

def _taskkill_by_names(names: List[str], force: bool = False) -> bool:
    if not names:
        return False
    ok = False
    for n in names:
        try:
            args = ["taskkill", "/IM", n, "/T"]
            if force:
                args.append("/F")
            rc = subprocess.call(args, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            ok = ok or (rc == 0)
        except Exception:
            pass
    return ok

# ---- no-tree taskkill (won't kill child processes) ----
def _taskkill_by_pids_no_tree(pids: List[int], force: bool = False) -> bool:
    if not pids:
        return False
    ok = False
    for pid in pids:
        try:
            args = ["taskkill", "/PID", str(pid)]  # NOTE: no /T
            if force:
                args.append("/F")
            rc = subprocess.call(args, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            ok = ok or (rc == 0)
        except Exception:
            pass
    return ok

def _taskkill_by_names_no_tree(names: List[str], force: bool = False) -> bool:
    if not names:
        return False
    ok = False
    for n in names:
        try:
            args = ["taskkill", "/IM", n]  # NOTE: no /T
            if force:
                args.append("/F")
            rc = subprocess.call(args, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            ok = ok or (rc == 0)
        except Exception:
            pass
    return ok

def _send_save_as_hotkey():
    try:
        if pyautogui:
            pyautogui.press('f12')
    except Exception:
        pass

def _desktop_path() -> str:
    return str(pathlib.Path.home() / "Desktop")

def word_save_active_as(full_path: str) -> bool:
    """
    Save the active Word document as 'full_path' (creates doc if needed).
    Returns True if saved, False otherwise.
    """
    if win32 is None:
        return False
    try:
        app = _word_app()
        if app is None:
            app = win32.gencache.EnsureDispatch('Word.Application')
            app.Visible = True
        if app.Documents.Count == 0:
            app.Documents.Add()
        # Ensure .docx extension if user forgot
        if not re.search(r"\.docx?$", full_path, re.I):
            full_path += ".docx"
        app.ActiveDocument.SaveAs2(full_path)
        return True
    except Exception as e:
        print(f"[WORD SAVE AS] {e}")
        return False

def _office_close_or_saveas(app_key: str) -> bool:
    if win32 is None:
        return False
    try:
        if app_key == "word":
            try:
                app = win32.GetActiveObject("Word.Application")
            except Exception:
                app = None
            if not app:
                return False
            app.Visible = True
            if app.Documents.Count == 0:
                app.Quit(); speak("Closed Word."); return True
            doc = app.ActiveDocument
            if doc.Saved:
                app.Quit(); speak("Closed Word.")
            else:
                try:
                    app.Dialogs(84).Show()
                except Exception:
                    _send_save_as_hotkey()
                speak("Opened Save As for Word. Pick a location and save, then say 'Hey Friday, close Word' again if needed.")
            return True

        if app_key == "excel":
            try:
                xl = win32.GetActiveObject("Excel.Application")
            except Exception:
                xl = None
            if not xl:
                return False
            xl.Visible = True
            if xl.Workbooks.Count == 0:
                xl.Quit(); speak("Closed Excel."); return True
            wb = xl.ActiveWorkbook
            if wb.Saved:
                xl.Quit(); speak("Closed Excel.")
            else:
                try:
                    xl.Application.Dialogs(5).Show()
                except Exception:
                    _send_save_as_hotkey()
                speak("Opened Save As for Excel. Choose where to save, then say 'Hey Friday, close Excel' again if needed.")
            return True

        if app_key == "powerpoint":
            try:
                ppt = win32.GetActiveObject("PowerPoint.Application")
            except Exception:
                ppt = None
            if not ppt:
                return False
            ppt.Visible = True
            if ppt.Presentations.Count == 0:
                ppt.Quit(); speak("Closed PowerPoint."); return True
            pres = ppt.ActivePresentation
            if pres.Saved:
                ppt.Quit(); speak("Closed PowerPoint.")
            else:
                _send_save_as_hotkey()
                speak("Opened Save As for PowerPoint. Choose where to save, then say 'Hey Friday, close PowerPoint' again if needed.")
            return True

        if app_key == "access":
            try:
                acc = win32.GetActiveObject("Access.Application")
            except Exception:
                acc = None
            if not acc:
                return False
            acc.Visible = True
            _send_save_as_hotkey()
            speak("Opened Save As in Access if there are changes. After saving, say 'Hey Friday, close Access' again.")
            return True

    except Exception:
        return False
    return False

# ---- shells that must not kill their children ----
_SHELL_KEYS = {"terminal", "powershell", "cmd", "wsl"}

def _close_shell_safely(target: str) -> bool:
    """
    Close only the terminal/shell process — never its child processes.
    Steps:
      1) Prefer recorded PIDs (from _OPENED_PIDS[target]).
      2) CloseMainWindow() on those PIDs.
      3) taskkill (no /T), then with /F if needed.
      4) If no PIDs, operate by process name, still no /T.
      (No Alt+F4 fallback here to avoid closing the foreground app by mistake.)
    """
    # 1) Try our recorded PIDs first
    pids = list(_OPENED_PIDS.get(target, set()))
    if pids:
        if _ps_close_main_window_by_pids(pids):
            speak(f"Asked {target} to close."); return True
        if _taskkill_by_pids_no_tree(pids, force=False):
            speak(f"Closed {target}."); return True
        if _taskkill_by_pids_no_tree(pids, force=True):
            speak(f"Force-closed {target}."); return True
        return False

    # 2) No recorded PIDs -> operate by process name, still no tree kills
    names = PROC_NAME_MAP.get(target, [])
    if _ps_close_main_window_by_names(names):
        speak(f"Asked {target} to close."); return True
    if _taskkill_by_names_no_tree(names, force=False):
        speak(f"Closed {target}."); return True
    if _taskkill_by_names_no_tree(names, force=True):
        speak(f"Force-closed {target}."); return True

    # 3) No matches found — do NOT send Alt+F4
    speak(f"{target.title()} doesn't seem to be running.")
    return False


def close_app(app: Optional[str]) -> bool:
    key = _normalize_app_name(app or "")
    target = key or (_last_opened_app_key or "")

    if not target:
        speak("Which app should I close?")
        return False

    # Shells: close without killing their child trees
    if target in _SHELL_KEYS:
        if _close_shell_safely(target):
            return True
        # If we couldn't find it, _close_shell_safely already spoke.
        return False

    # Office apps: handle Save As if there are unsaved changes
    if target in ("word", "excel", "powerpoint", "access"):
        handled = _office_close_or_saveas(target)
        if handled:
            return True
        # If Office COM isn't open, fall through to name-based close.

    # Generic path (OK to use /T here for non-shells)
    pids = list(_OPENED_PIDS.get(target, set()))
    if _ps_close_main_window_by_pids(pids):
        speak(f"Asked {target} to close."); return True
    if _taskkill_by_pids(pids, force=False):
        speak(f"Closed {target}."); return True
    if _taskkill_by_pids(pids, force=True):
        speak(f"Force-closed {target}."); return True

    names = PROC_NAME_MAP.get(target, [])
    if _ps_close_main_window_by_names(names):
        speak(f"Asked {target} to close."); return True
    if _taskkill_by_names(names, force=False):
        speak(f"Closed {target}."); return True
    if _taskkill_by_names(names, force=True):
        speak(f"Force-closed {target}."); return True

    # IMPORTANT: Do NOT Alt+F4 — we don't want to close the active window by mistake
    if names:
        speak(f"{target.title()} doesn't seem to be running.")
    else:
        speak(f"I couldn't find {target}.")
    return False

# ============================ Intent & heuristics ============================
def _wants_browser(low: str) -> bool:

    low = " " + low.lower() + " "
    triggers = [
        " in browser ", " on browser ", " open in browser ",
        " on google ", " on the web ", " on web ",
        " open google ", " open a tab ", " show me on google ",
        " show in browser ", " open the browser "
    ]
    return any(t in low for t in triggers)

_TIME_FILLERS = r"(?:right\s*now|now|currently|at\s+the\s+moment|rn)"
def _extract_time_intent(low: str) -> Optional[dict]:
    want_browser = _wants_browser(low)
    low = low.strip()
    city_patterns = [
        rf"(?:what(?:'s| is)?\s+)?the\s*time\s+in\s+([a-z\s,]+?)(?:\s+{_TIME_FILLERS})?\s*\??$",
        rf"what\s+time\s+is\s+it\s+in\s+([a-z\s,]+?)(?:\s+{_TIME_FILLERS})?\s*\??$",
        rf"(?:current\s+time|time)\s+in\s+([a-z\s,]+?)\s*\??$",
        rf".*?\btime\s+in\s+([a-z\s,]+?)\s*(?:{_TIME_FILLERS})?\s*\??$",
    ]
    for pat in city_patterns:
        m = re.match(pat, low) or re.search(pat, low)
        if m:
            city = (m.group(1) or "").strip()
            city = re.sub(rf"\b{_TIME_FILLERS}\b", "", city).strip()
            if city:
                return {"intent": "get_time", "args": {"city": city, "open": want_browser}}

    local_patterns = [
        rf"^(?:what(?:'s| is)?\s+)?the\s*time(?:\s+{_TIME_FILLERS})?\s*\??$",
        rf"^what\s+time\s+is\s+it(?:\s+{_TIME_FILLERS})?\s*\??$",
        r"^(?:current\s+time|time)\s*(?:please)?\s*\??$",
        r"^tell\s+me\s+the\s+time\s*\??$",
        r"^give\s+me\s+the\s+time\s*\??$",
    ]
    for pat in local_patterns:
        if re.match(pat, low):
            return {"intent": "get_local_time", "args": {"open": want_browser}}
    if "time" in low and " in " not in low:
        return {"intent": "get_local_time", "args": {"open": want_browser}}
    return None

_WEATHER_FILLERS = r"(?:right\s*now|now|today|currently)"
def _extract_weather_intent(low: str) -> Optional[dict]:
    want_browser = _wants_browser(low)
    pats = [
        rf"(?:what(?:'s| is)?\s+)?the\s+weather\s+in\s+([a-z\s,]+?)(?:\s+{_WEATHER_FILLERS})?\s*\??$",
        rf"(?:weather|forecast)\s+in\s+([a-z\s,]+?)\s*\??$",
        rf".*?\bweather\s+in\s+([a-z\s,]+?)\b.*$",
    ]
    for pat in pats:
        m = re.match(pat, low) or re.search(pat, low)
        if m:
            city = (m.group(1) or "").strip()
            if city:
                return {"intent": "get_weather", "args": {"city": city, "open": want_browser}}
    return None

def _extract_news_intent(low: str) -> Optional[dict]:
    want_browser = _wants_browser(low)
    pats = [
        rf"(?:what(?:'s| is)?\s+)?the\s+news\s+in\s+([a-z\s,]+?)\s*\??$",
        rf"(?:news|headlines)\s+in\s+([a-z\s,]+?)\s*\??$",
        rf".*?\bnews\s+in\s+([a-z\s,]+?)\b.*$",
    ]
    for pat in pats:
        m = re.match(pat, low) or re.search(pat, low)
        if m:
            city = (m.group(1) or "").strip()
            if city:
                return {"intent": "get_news", "args": {"city": city, "open": want_browser}}
    return None

def extract_intent(utterance: str) -> dict:
    low = utterance.lower().strip()
    low = re.sub(r'^[\s\.,!?;:~\-]+', '', low)

    # --- Multi-action detection: if user chained commands with commas / and / then ---
    if ("," in low or " and " in low or " then " in low) and len(low.split()) >= 6:
        plan = _planner_llm(utterance) or _fallback_plan(utterance)
        if plan and isinstance(plan.get("steps"), list) and plan["steps"]:
            return {"intent": "multi_action", "args": {"steps": plan["steps"]}}

    if re.search(r"\bstop recording meeting\b", low):
        return {"intent": "meeting_record_stop", "args": {}}
    if re.search(r"\bstop recording\b", low):
        return {"intent": "meeting_record_stop", "args": {}}
    if re.search(r"\bstop meeting\b", low):
        return {"intent": "meeting_record_stop", "args": {}}
    if re.search(r"\bend meeting\b", low):
        return {"intent": "meeting_record_stop", "args": {}}
    if re.search(r"\bend recording\b", low):
        return {"intent": "meeting_record_stop", "args": {}}
    if re.search(r"\bend recording meeting\b", low):
        return {"intent": "meeting_record_stop", "args": {}}

    if re.search(r"\brecord meeting\b", low):
        return {"intent": "meeting_record_start", "args": {}}
    if re.search(r"\bstart meeting\b", low):
        return {"intent": "meeting_record_start", "args": {}}
    if re.search(r"\bstart meeting recording\b", low):
        return {"intent": "meeting_record_start", "args": {}}
    if re.search(r"\bstart recording meeting\b", low):
        return {"intent": "meeting_record_start", "args": {}}

    m_close1 = re.match(r"^(?:close|quit|exit)\s+(.*)$", low)
    if m_close1:
        candidate = m_close1.group(1).strip(" .!?")
        candidate = re.sub(r"\b(app(?:lication)?|window|the app)\b", "", candidate).strip()
        return {"intent": "close_app", "args": {"app": candidate or ""}}
    if re.match(r"^(?:close|quit|exit)\s*(?:the\s+)?(?:app|application|window)?\s*$", low):
        return {"intent": "close_app", "args": {"app": ""}}

    t = _extract_time_intent(low)
    if t:
        return t

    w = _extract_weather_intent(low)
    if w:
        return w
    n = _extract_news_intent(low)
    if n:
        return n

    m2 = re.search(r"open\s+browser\s+(?:and\s+)?search\s+for\s+(.+)", low)
    if m2:
        q = m2.group(1).strip(" .!?")
        new_tab = "new tab" in low
        return {"intent": "browser_search", "args": {"query": q, "new_tab": new_tab}}

    try:
        content = call_llm(
            [{"role": "user", "content": utterance}],
            system=INTENT_SYSTEM,
            temperature=0.2,
            max_tokens=256
        )
        j = re.search(r"\{.*\}", content, re.S)
        if j:
            return json.loads(j.group(0))
    except Exception:
        pass

    low = utterance.lower().strip().rstrip(".!?")
    if low.startswith("open powershell"):
        return {"intent": "open_powershell", "args": {}}
    if low.startswith("open terminal"):
        return {"intent": "open_terminal", "args": {}}
    if low.startswith("open cmd") or low.startswith("open command prompt"):
        return {"intent": "open_cmd", "args": {}}
    if low.startswith("open wsl"):
        return {"intent": "open_wsl", "args": {}}
    if low.startswith("open "):
        app = utterance.split(" ", 1)[1].strip().rstrip(".!?")
        return {"intent": "open_app", "args": {"app": app}}
    if low.startswith("search "):
        return {"intent": "browser_search", "args": {"query": low.split(" ", 1)[1], "new_tab": False}}
    if "write" in low and "essay" in low:
        return {"intent": "write_essay", "args": {"topic": utterance, "length_words": _extract_words_hint(utterance)}}

    # open document/file at ...
    m = re.search(r'\bopen\s+(?:the\s+)?(?:document|file)\s+(?:at\s+)?(.+)$', low)
    if m:
        return {"intent": "open_document", "args": {"path": m.group(1).strip()}}

    # edit/update/add/delete in document ...
    if low.startswith("edit document") or low.startswith("update document") or low.startswith("modify document") \
    or low.startswith("add ") or low.startswith("append ") or low.startswith("insert ") \
    or low.startswith("delete ") or low.startswith("remove ") or low.startswith("replace "):
        return {"intent": "edit_document", "args": {"instruction": utterance.strip()}}

    return {"intent": "ask", "args": {"question": utterance}}


# ============================ Safety confirm ============================
def manual_confirm_dialog(message: str) -> bool:
    try:
        import ctypes
        MB_YESNO = 0x00000004
        MB_ICONWARNING = 0x00000030
        MB_TOPMOST = 0x00040000
        res = ctypes.windll.user32.MessageBoxW(0, message, "Confirm", MB_YESNO | MB_ICONWARNING | MB_TOPMOST)
        return res == 6
    except Exception:
        try:
            print(f"[CONFIRM] {message}  Type YES to confirm: ", end="", flush=True)
            resp = input().strip()
            return resp.upper() == "YES"
        except Exception:
            return False

def confirm_before_download(url: str) -> bool:
    return manual_confirm_dialog(f"Download from:\n{url}\n\nProceed?")


# ============================ App normalization & Office openers ============================
def _normalize_app_name(app: str) -> str:
    if not app:
        return ""
    app = app.strip().lower()
    app = app.strip(string.punctuation + " ")
    app = re.sub(r"\s+", " ", app)
    for prefix in ["microsoft ", "ms "]:
        if app.startswith(prefix):
            app = app[len(prefix):]
    app = app.rstrip(".")
    if app in ("ms", "app", "apps", "application", "applications", "window", "the app"):
        return ""
    synonyms = {
        "word": "word", "winword": "word", "ms word": "word",
        "excel": "excel", "ms excel": "excel",
        "powerpoint": "powerpoint", "ppt": "powerpoint", "ms powerpoint": "powerpoint",
        "access": "access", "ms access": "access",
        "gmail": "gmail", "outlook": "outlook",
        "google calendar": "google calendar", "calendar google": "google calendar",
        "visual studio code": "vscode", "vs code": "vscode", "code": "vscode",
        "command prompt": "cmd",
        "windows terminal": "terminal", "terminal": "terminal",
        "powershell": "powershell",
        "edge": "edge",
        "chrome": "chrome", "google chrome": "chrome",
        "notepad": "notepad",
        "spotify": "spotify",
        "calculator": "calculator", "calc": "calculator",
        "file explorer": "explorer", "explorer": "explorer",
        "calendar": "calendar",
        "browser": "browser",
    }
    return synonyms.get(app, app)

def _first_existing(paths: List[str]) -> Optional[str]:
    for p in paths:
        pe = os.path.expandvars(p)
        if os.path.exists(pe):
            return pe
    return None

def _fallback_ctrl_n(delay: float = 1.0):
    try:
        if pyautogui:
            time.sleep(delay)
            pyautogui.hotkey('ctrl', 'n')
    except Exception:
        pass

def _word_app():
    if win32 is None:
        return None
    try:
        return win32.GetActiveObject("Word.Application")
    except Exception:
        try:
            return win32.gencache.EnsureDispatch('Word.Application')
        except Exception:
            return None

def _open_word_new():
    if win32 is not None:
        app = _word_app()
        if app:
            app.Visible = True
            app.Documents.Add()
            speak("Opened a new Word document.")
            _record_opened_pid("word", None)
            return True
    exe = _first_existing(APP_PATHS["word_paths"])
    try:
        p = subprocess.Popen('start "" winword', shell=True) if exe is None else subprocess.Popen(f'"{exe}"', shell=True)
        speak("Opening Word.")
        _record_opened_pid("word", getattr(p, "pid", None))
        _fallback_ctrl_n(1.2)
        return True
    except Exception:
        return False

def _open_excel_new():
    if win32 is not None:
        try:
            xl = None
            try:
                xl = win32.GetActiveObject("Excel.Application")
            except Exception:
                xl = win32.gencache.EnsureDispatch('Excel.Application')
            xl.Visible = True
            xl.Workbooks.Add()
            speak("Opened a new Excel workbook.")
            _record_opened_pid("excel", None)
            return True
        except Exception:
            pass
    exe = _first_existing(APP_PATHS["excel_paths"])
    try:
        p = subprocess.Popen('start "" excel', shell=True) if exe is None else subprocess.Popen(f'"{exe}"', shell=True)
        speak("Opening Excel.")
        _record_opened_pid("excel", getattr(p, "pid", None))
        _fallback_ctrl_n(1.2)
        return True
    except Exception:
        return False

def _open_powerpoint_new():
    if win32 is not None:
        try:
            ppt = None
            try:
                ppt = win32.GetActiveObject("PowerPoint.Application")
            except Exception:
                ppt = win32.gencache.EnsureDispatch('PowerPoint.Application')
            ppt.Visible = True
            ppt.Presentations.Add()
            speak("Opened a new PowerPoint presentation.")
            _record_opened_pid("powerpoint", None)
            return True
        except Exception:
            pass
    exe = _first_existing(APP_PATHS["powerpoint_paths"])
    try:
        if exe:
            p = subprocess.Popen(f'"{exe}" /N', shell=True)
        else:
            p = subprocess.Popen('start "" powerpnt /N', shell=True)
        speak("Opening PowerPoint.")
        _record_opened_pid("powerpoint", getattr(p, "pid", None))
        _fallback_ctrl_n(1.4)
        return True
    except Exception:
        return False

def _open_access_new():
    if win32 is not None:
        try:
            acc = None
            try:
                acc = win32.GetActiveObject("Access.Application")
            except Exception:
                acc = win32.gencache.EnsureDispatch('Access.Application')
            acc.Visible = True
            tmp_dir = os.path.expandvars(r"%TEMP%"); os.makedirs(tmp_dir, exist_ok=True)
            db_path = os.path.join(tmp_dir, f"NewDatabase_{int(time.time())}.accdb")
            acc.NewCurrentDatabase(db_path)
            speak("Opened a new Access database.")
            _record_opened_pid("access", None)
            return True
        except Exception:
            pass
    exe = _first_existing(APP_PATHS["access_paths"])
    try:
        if exe:
            p = subprocess.Popen(f'"{exe}"', shell=True)
        else:
            p = subprocess.Popen('start "" msaccess', shell=True)
        speak("Opening Access.")
        _record_opened_pid("access", getattr(p, "pid", None))
        try:
            if pyautogui:
                time.sleep(1.5); pyautogui.hotkey('ctrl', 'n'); time.sleep(0.5); pyautogui.press('enter')
        except Exception:
            pass
        return True
    except Exception:
        return False

OFFICE_OPENERS = {
    "word": _open_word_new,
    "excel": _open_excel_new,
    "powerpoint": _open_powerpoint_new,
    "access": _open_access_new,
}

def _open_system_app(key: str) -> bool:
    k = key.lower()
    try:
        if k == "calculator":
            p = subprocess.Popen("calc", shell=True); speak("Opening Calculator."); _record_opened_pid("calculator", getattr(p, "pid", None)); return True
        if k == "explorer":
            p = subprocess.Popen("explorer", shell=True); speak("Opening File Explorer."); _record_opened_pid("explorer", getattr(p, "pid", None)); return True
        if k == "calendar":
            if subprocess.call("start outlookcal:", shell=True) == 0:
                speak("Opening Calendar."); _record_opened_pid("calendar", None); return True
            if subprocess.call("start ms-calendar:", shell=True) == 0:
                speak("Opening Calendar."); _record_opened_pid("calendar", None); return True
            _browser.open_url("https://calendar.google.com/"); speak("Opening Google Calendar."); _record_opened_pid("browser", None); return True
        if k == "browser":
            _browser.open_url("https://www.google.com/"); speak("Opening browser."); _record_opened_pid("browser", None); return True
    except Exception:
        pass
    return False

def _launch_alias_or_path(key: str) -> bool:
    key = _normalize_app_name(key)

    if open_web_shortcut(key):
        return True

    if key in OFFICE_OPENERS:
        ok = OFFICE_OPENERS[key]()
        if ok:
            _record_opened_pid(key, None)
        return ok

    if key in ("calculator", "explorer", "calendar", "browser"):
        if _open_system_app(key):
            return True

    aliases = APP_ALIASES.get(key, []) + [key]
    for alias in aliases:
        exe = shutil.which(alias)
        if exe:
            try:
                p = subprocess.Popen(exe, shell=False)
                speak(f"Opening {alias}.")
                _record_opened_pid(key, getattr(p, "pid", None))
                return True
            except Exception:
                pass

    if key in APP_PATHS:
        path = os.path.expandvars(APP_PATHS[key])
        if os.path.exists(path):
            try:
                p = subprocess.Popen(f'"{path}"', shell=True)
                speak(f"Opening {key}.")
                _record_opened_pid(key, getattr(p, "pid", None))
                return True
            except Exception:
                pass

    low = key.lower()
    for canon, alias_list in APP_ALIASES.items():
        for alias in alias_list:
            if low in alias.lower() or alias.lower() in low:
                exe = shutil.which(alias)
                if exe:
                    try:
                        p = subprocess.Popen(exe, shell=False)
                        speak(f"Opening {alias}.")
                        _record_opened_pid(key, getattr(p, "pid", None))
                        return True
                    except Exception:
                        pass

    return False

def open_app(app: str) -> bool:
    global _awaiting_app_name
    app = _normalize_app_name(app)
    if not app:
        _awaiting_app_name = True
        speak("Which app?")
        return False
    _awaiting_app_name = False
    if _launch_alias_or_path(app):
        return True
    speak(f"I couldn't find {app}.")
    return False

def open_terminal():   _launch_alias_or_path("terminal") or _launch_alias_or_path("powershell")
def open_powershell(): _launch_alias_or_path("powershell")
def open_cmd():        _launch_alias_or_path("cmd")
def open_wsl():        _launch_alias_or_path("wsl")

def open_browser(url: Optional[str] = None):
    _browser.open_url(url or "https://www.google.com/")
    speak("Browser opened.")
    _record_opened_pid("browser", None)

def browser_search(query: str, new_tab: bool = False):
    _browser.search(query, new_tab=new_tab)
    speak(f"Searching for {query}.")
    _record_opened_pid("browser", None)

def web_search(query: str):
    browser_search(query, new_tab=False)


# ============================ Recycle bin & deletion ============================
def _recycle_bin_restore(name_contains: Optional[str] = None):
    name_filter = ""
    if name_contains:
        needle = name_contains.replace("'", "''")
        name_filter = f" | Where-Object {{ $_.Name -like '*{needle}*' -or $_.OriginalPath -like '*{needle}*' }}"
    ps = "try { Get-ChildItem -Path 'Recycle:\\' " + name_filter + " | Restore-RecycleBin -Force } catch { }"
    _ps_run(ps)

def restore_recycle_bin(name_contains: Optional[str] = None):
    try:
        _recycle_bin_restore(name_contains)
        speak("Attempted restore from Recycle Bin.")
    except Exception as e:
        speak(f"Restore failed: {e}")

def _rmdir_permanent(path: str):
    if os.path.isdir(path):
        subprocess.call(['cmd', '/c', 'rmdir', '/s', '/q', path])
    else:
        try:
            os.remove(path)
        except IsADirectoryError:
            subprocess.call(['cmd', '/c', 'rmdir', '/s', '/q', path])

def clear_folder(path: str, permanent: bool = False):
    if not os.path.isdir(path):
        speak("That path isn't a folder.")
        return
    if not manual_confirm_dialog(
        (f"PERMANENTLY " if permanent else "") +
        f"Delete ALL items INSIDE:\n{path}\n\n"
        "Click Yes to confirm. (Voice confirmations are ignored.)"
    ):
        speak("Canceled."); return

    if not permanent:
        if send2trash is None:
            speak("Install send2trash for safe deletes: pip install send2trash"); return
        for name in os.listdir(path):
            full = os.path.join(path, name)
            try:
                send2trash(full)
            except Exception as e:
                print(f"[send2trash] {full}: {e}")
        speak("Folder cleared to Recycle Bin.")
    else:
        for name in os.listdir(path):
            full = os.path.join(path, name)
            try:
                _rmdir_permanent(full)
            except Exception as e:
                print(f"[delete] {full}: {e}")
        speak("Folder permanently cleared.")

def delete_path(path: str, permanent: bool = False):
    if not os.path.exists(path):
        speak("That path doesn't exist."); return
    if not manual_confirm_dialog(
        (f"PERMANENTLY " if permanent else "") +
        f"Delete this item:\n{path}\n\n"
        "Click Yes to confirm. (Voice confirmations are ignored.)"
    ):
        speak("Canceled."); return

    if not permanent and send2trash is not None:
        try:
            send2trash(path); speak("Item moved to Recycle Bin."); return
        except Exception as e:
            print(f"[send2trash] {path}: {e}")

    try:
        _rmdir_permanent(path); speak("Item permanently deleted.")
    except Exception as e:
        speak(f"Delete failed: {e}")


# ============================ Writer & coder ============================
def _topic_subject(t: str) -> str:
    t = t.strip()
    m = re.search(r"\babout\s+(.+)", t, re.I)
    if m:
        return m.group(1).strip().rstrip(".")
    # strip leading “write …”, “draft …” etc.
    t2 = re.sub(r"^(write|draft|create|make)\b.*?\babout\b", "", t, flags=re.I).strip()
    return (t2 or t).rstrip(".")


def _fallback_essay(topic: str, n_words: int) -> str:
    subject = _topic_subject(topic) or "the topic"
    base = [
        f"{subject.title()} matters because it blends practical uses with fast-evolving ideas.",
        "At a high level, it helps people save time, reduce errors, and uncover patterns they would otherwise miss.",
        "In everyday work it shows up in tools that draft text, summarize notes, analyze data, and answer questions.",
        "Good results still depend on clear goals, well-chosen data, and human judgment.",
        "Looking ahead, expect better accuracy, safer behavior, and easier ways to customize models for specific tasks."
    ]
    text = " ".join(base)

    # Deterministic padding so we can hit the target without an LLM
    filler_tpls = [
        f"For example, in {subject}, teams often balance accuracy, speed, and cost while iterating in small steps.",
        "Clear evaluation, simple baselines, and feedback loops usually improve results.",
        "Real-world constraints such as privacy, reliability, and maintenance shape practical choices.",
        "When possible, start small, measure, and expand only after you see consistent gains."
    ]
    i = 0
    while len(re.findall(r"\b\w+\b", text)) < n_words:
        text += " " + filler_tpls[i % len(filler_tpls)]
        i += 1

    # Trim to clean sentence end if we overshoot
    words = text.split()
    if len(words) > n_words:
        text = " ".join(words[:n_words]) + "."

    return text.strip()

def _word_count(s: str) -> int:
    return len(re.findall(r"\b\w+\b", s))

def _pad_to_min_words(text: str, target: int, topic: str) -> str:
    # Non-LLM local padding, used if the model call failed or was short.
    if _word_count(text) >= target:
        return text
    extra_bits = [
        f"In practice, applying {topic} benefits from clear objectives and small experiments.",
        "Adding examples, measuring outcomes, and refining assumptions generally leads to steady improvement.",
    ]
    i = 0
    while _word_count(text) < target:
        text += " " + extra_bits[i % len(extra_bits)]
        i += 1
    return text

def generate_text(prompt: str, words: int = 300) -> str:
    target = max(50, min(800, words))
    sys_prompt = (
        "You are a helpful, concise writer. "
        "Write plain paragraphs (no headings, no bullets). "
        f"Produce AT LEAST {int(target*0.95)} words and ABOUT {target} words total."
    )
    user = f"Topic: {prompt}\n\nWrite an essay."

    out = None
    try:
        out = call_llm(
            [{"role": "user", "content": user}],
            system=sys_prompt,
            temperature=0.6,
            max_tokens=max(512, target * 3)
        ).strip()
    except Exception as e:
        speak(f"LLM unavailable: {e}")
        out = _fallback_essay(prompt, target)

    # If LLM responded but is short, first try an LLM expansion
    if out and _word_count(out) < int(target*0.9):
        try:
            need = max(0, target - _word_count(out))
            expand_prompt = (
                f"Expand the same essay by at least {need} more words. "
                "Add detail and examples, keep the same topic and tone. "
                "Return only the additional paragraphs."
            )
            extra = call_llm(
                [{"role": "user", "content": expand_prompt}],
                system="You add paragraphs that continue the essay.",
                temperature=0.6,
                max_tokens=max(256, need * 3)
            ).strip()
            out = (out + "\n\n" + extra).strip()
        except Exception:
            # If expansion also fails, pad locally
            out = _pad_to_min_words(out, target, prompt)

    # Final guard: if still short (e.g., LLM totally down), pad locally
    if _word_count(out) < int(target*0.95):
        out = _pad_to_min_words(out, target, prompt)

    return out
    

def write_to_word(text: str) -> bool:
    if win32 is None:
        return False
    app = _word_app()
    if app is None:
        try:
            app = win32.gencache.EnsureDispatch('Word.Application')
            app.Visible = True
        except Exception:
            return False
    try:
        if app.Documents.Count == 0:
            app.Documents.Add()
        app.Selection.EndKey(Unit=6)
        app.Selection.TypeParagraph()
        app.Selection.TypeText(text)
        return True
    except Exception as e:
        print(f"[WORD] {e}")
    return False


def write_to_notepad(text: str) -> bool:
    try:
        p = subprocess.Popen('notepad', shell=True)
        _record_opened_pid("notepad", getattr(p, "pid", None))
        time.sleep(0.8)
        pyperclip.copy(text)
        if pyautogui:
            pyautogui.hotkey('ctrl', 'v')
            return True
        return False
    except Exception as e:
        print(f"[Notepad] {e}")
        return False


def do_write_essay(topic: str, length_words: Optional[int] = None):
    words = max(50, min(800, int(length_words) if length_words else 150))
    speak(f"Writing an essay on {topic}.")
    text = generate_text(topic, words)
    wc = len(re.findall(r"\b\w+\b", text))
    print(f"[DEBUG] Generated {wc} words (target {words}).")
    if write_to_word(text):
        speak("I wrote it in Word.")
    elif write_to_notepad(text):
        speak("Word wasn't available, so I pasted it into Notepad.")
    else:
        speak("Couldn't open an editor; printing in console.")
        print("\n===== ESSAY =====\n" + text + "\n==================\n")


def _extract_words_hint(text: str, *, default=150, lo=50, hi=800) -> int:
    low = text.lower()
    m = re.search(r"\b(\d{2,4})\b", low)
    if m:
        n = int(m.group(1))
        return max(lo, min(hi, n))
    words_map = {
        "fifty":50, "sixty":60, "seventy":70, "eighty":80, "ninety":90,
        "hundred":100, "two hundred":200, "three hundred":300, "four hundred":400,
        "five hundred":500, "six hundred":600, "seven hundred":700, "eight hundred":800
    }
    for k, v in words_map.items():
        if k in low:
            return max(lo, min(hi, v))
    return default



def generate_code(language: Optional[str], prompt: str, open_in: Optional[str] = "vscode"):
    language = (language or "python").lower()
    sys_prompt = (
        "You are a senior engineer. Generate a single self-contained code file. "
        "Be readable; include minimal comments."
    )
    code = call_llm([
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": f"Write a {language} program: {prompt}"},
    ])
    m = re.search(r"```[a-zA-Z0-9]*\n([\s\S]*?)```", code)
    body = m.group(1).strip() if m else code.strip()

    ext = {
        'python': 'py','py':'py','cpp':'cpp','c++':'cpp','c':'c','javascript':'js','js':'js','ts':'ts',
        'java':'java','go':'go','rust':'rs','html':'html','css':'css'
    }.get(language, 'txt')
    out_path = os.path.abspath(f"generated.{ext}")
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(body)
    speak(f"Code written to {out_path}.")
    if open_in == "vscode":
        _launch_alias_or_path("vscode")
        time.sleep(0.8)
        try:
            subprocess.Popen(f'code "{out_path}"', shell=True)
        except Exception:
            pass
    elif open_in == "notepad":
        subprocess.Popen(f'notepad "{out_path}"', shell=True)

# ============================ Spotify & media keys ============================
def spotify_control(action: str, playlist_url: Optional[str] = None):
    act = (action or "").lower()
    if act in ("play", "pause", "playpause") and keyboard:
        keyboard.send("play/pause media"); speak("Toggled play/pause."); return
    if act == "next" and keyboard:
        keyboard.send("next track"); speak("Next track."); return
    if act in ("prev", "previous") and keyboard:
        keyboard.send("previous track"); speak("Previous track."); return

    try:
        import spotipy
        from spotipy.oauth2 import SpotifyOAuth
        scope = "user-modify-playback-state user-read-playback-state"
        sp = spotipy.Spotify(auth_manager=SpotifyOAuth(scope=scope))
        if act == "play_playlist" and playlist_url:
            sp.start_playback(context_uri=playlist_url); speak("Playing playlist."); return
        if act == "play":
            sp.start_playback(); speak("Playing."); return
        if act == "pause":
            sp.pause_playback(); speak("Paused."); return
        if act == "next":
            sp.next_track(); speak("Next track."); return
        if act in ("prev", "previous"):
            sp.previous_track(); speak("Previous track."); return
    except Exception as e:
        speak(f"Spotify control not available: {e}")

# ============================ Generic Q&A ============================
def do_ask(question: str):
    try:
        answer = call_llm(
            [{"role": "user", "content": question}],
            system="You are a helpful assistant. Answer briefly and clearly.",
            temperature=0.5,
            max_tokens=4096
        )
        speak(answer)
    except Exception as e:
        speak(f"LLM error: {e}")

# ============================ Wake phrase & loop ============================
def extract_after_wake(text: str) -> Optional[str]:
    if not text:
        return None
    
    low = text.lower()
    for w in WAKE_PHRASES:
        m = re.search(r'\b' + re.escape(w) + r'\b', low)
        if m:
            remainder = text[m.end():]

            remainder = re.sub(r'^[\s\.,!?;:~\-]+', '', remainder)
            return remainder.strip()
    return None


def handle_utterance(raw_text: str):
    global _awaiting_app_name


    if _awaiting_app_name:
        candidate = raw_text.strip()
        low = re.sub(r"[^\w\s]", "", candidate.lower()).strip()
        for prefix in ("open ", "launch ", "start "):
            if low.startswith(prefix):
                candidate = candidate[len(prefix):].strip()
                low = re.sub(r"[^\w\s]", "", candidate.lower()).strip()
                break

        if low in ("cancel", "nevermind", "never mind", "stop"):
            _awaiting_app_name = False
            speak("Okay, canceled.")
            return

        app_norm = _normalize_app_name(candidate)
        if not app_norm:
            speak("I didn't catch the app name. For example, Word or Excel.")
            return
        open_app(app_norm)
        return


    cmd = extract_after_wake(raw_text)
    norm = re.sub(r"[^\w\s]", "", (cmd or "")).lower().strip()
    if norm in ("stop", "exit", "quit"):
        speak("Stopping. Goodbye.")
        raise SystemExit
    if cmd is None or cmd == "":
        return
    if cmd.lower().strip() in ("stop", "exit", "quit"):
        speak("Stopping. Goodbye.")
        raise SystemExit

    intent = extract_intent(cmd)
    print(f"[🔧] Intent → {intent}")
    name = (intent or {}).get("intent")
    args = (intent or {}).get("args", {})

    if name == "meeting_record_start":
        if meeting_rec.running:
            speak("I'm already recording your meeting.")
        else:
            ok = meeting_rec.start()
            if ok:
                speak("Recording started.")
            else:
                speak("I couldn't start the recording.")
        return

    if name == "meeting_record_stop":
        if not meeting_rec.running:
            speak("There isn't an active meeting recording.")
            return
        audio_path = meeting_rec.stop()
        speak("Recording stopped. Transcribing now.")
        try:
            segments, _info = whisper.transcribe(audio_path, vad_filter=True, beam_size=5)
            text_lines = []
            for seg in segments:
                t = (seg.text or "").strip()
                if t:
                    text_lines.append(t)
            transcript = "\n".join(text_lines).strip()
            meeting_rec.write_docx(transcript)
            a = os.path.basename(meeting_rec.audio_path or "")
            d = os.path.basename(meeting_rec.docx_path or "")
            speak(f"Saved {a} and {d}.")
        except Exception as e:
            speak(f"Sorry, I couldn't transcribe that. {e}")
        return


    if name == "get_local_time":
        speak(f"The current time is {get_local_time_str()}.")
    elif name == "get_time":
        city = args.get("city", "")
        open_flag = bool(args.get("open", False))
        open_city_time(city, open_in_browser=open_flag, new_tab=False)
    elif name == "get_weather":
        city = args.get("city", "")
        open_flag = bool(args.get("open", False))
        open_city_weather(city, open_in_browser=open_flag, new_tab=False)
    elif name == "get_news":
        city = args.get("city", "")
        open_flag = bool(args.get("open", False))
        open_city_news(city, open_in_browser=open_flag, new_tab=False)
    elif name == "browser_search":
        browser_search(args.get("query", ""), bool(args.get("new_tab", False)))
    elif name == "open_app":
        open_app(args.get("app", ""))
    elif name == "close_app":
        close_app(args.get("app", ""))
    elif name == "open_terminal":
        open_terminal()
    elif name == "open_powershell":
        open_powershell()
    elif name == "open_cmd":
        open_cmd()
    elif name == "open_wsl":
        open_wsl()
    elif name == "open_browser":
        open_browser(args.get("url"))
    elif name == "web_search":
        web_search(args.get("query", ""))
    elif name == "clear_folder":
        clear_folder(args.get("path", ""), bool(args.get("permanent", False)))
    elif name == "delete_path":
        delete_path(args.get("path", ""), bool(args.get("permanent", False)))
    elif name == "restore_recycle_bin":
        restore_recycle_bin(args.get("name_contains"))
    elif name == "write_essay":
        do_write_essay(args.get("topic", ""), args.get("length_words"))
    elif name == "generate_code":
        generate_code(args.get("language"), args.get("prompt", ""), args.get("open_in", "vscode"))
    elif name == "spotify_control":
        spotify_control(args.get("action", ""), args.get("playlist_url"))
    elif name == "ask":
        do_ask(args.get("question", cmd))
    elif name == "open_document":
        open_document(args.get("path",""))
    elif name == "edit_document":
        edit_active_document(args.get("instruction",""))
    elif name == "multi_action":
        steps = args.get("steps", [])
        if not steps:
            speak("I couldn't plan the steps.")
        else:
            # 🔧 inject save step if the utterance contained a save request
            steps = _augment_plan_with_save(cmd, steps)
            speak("Okay, working through your steps.")
            execute_plan(steps)

    else:
        speak("I didn't understand that.")


def main():
    speak("Friday online. Say 'Hey Friday' followed by your request.")
    tmp_wav = "_last.wav"
    try:
        while True:
            try:
                # If muted from the bubble, idle here
                if _muted:
                    time.sleep(0.3)
                    continue

                record_wav_tmp(tmp_wav, CHUNK_SEC)
                text = transcribe(tmp_wav)
                handle_utterance(text)
                # fall back to listening state after handling

            except KeyboardInterrupt:
                speak("Goodbye.")
                break
            except SystemExit:
                break
            except Exception as e:
                print(f"[loop] error: {e}")
                time.sleep(0.2)
    finally:
        try:
            _browser.close()
        except Exception:
            pass

if __name__ == "__main__":
    main()
